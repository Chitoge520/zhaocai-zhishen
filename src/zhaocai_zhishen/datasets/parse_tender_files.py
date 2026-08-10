from __future__ import annotations

import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document


FIELD_PATTERNS = {
    "project_name": [
        r"项目名称[：:\s]+([^\n。；;]{4,120})",
        r"采购项目名称[：:\s]+([^\n。；;]{4,120})",
    ],
    "project_no": [
        r"项目编号[：:\s]+([A-Za-z0-9][A-Za-z0-9\-_/（）()【】\[\].]{2,})",
        r"采购编号[：:\s]+([A-Za-z0-9][A-Za-z0-9\-_/（）()【】\[\].]{2,})",
        r"招标编号[：:\s]+([A-Za-z0-9][A-Za-z0-9\-_/（）()【】\[\].]{2,})",
    ],
    "purchaser": [
        r"采购人[：:\s]+([^\n。；;]{2,80})",
        r"招标人[：:\s]+([^\n。；;]{2,80})",
    ],
    "budget_amount": [
        r"预算金额[：:\s]*([0-9,.，]+)\s*万元",
        r"采购预算[：:\s]*([0-9,.，]+)\s*万元",
        r"预算金额[：:\s]*([0-9,.，]+)\s*元",
        r"采购预算[：:\s]*([0-9,.，]+)\s*元",
        r"采购包预算金额[：:\s]*([0-9,.，]+)\s*元",
    ],
    "ceiling_price": [
        r"最高限价[：:\s]*([0-9,.，]+)\s*万元",
        r"最高投标限价[：:\s]*([0-9,.，]+)\s*万元",
        r"最高限价[：:\s]*([0-9,.，]+)\s*元",
        r"最高投标限价[：:\s]*([0-9,.，]+)\s*元",
    ],
    "ceiling_rate": [
        r"最高限价[：:\s]*[^0-9\n]{0,20}([0-9]+(?:\.[0-9]+)?)\s*%",
        r"最高投标限价[：:\s]*[^0-9\n]{0,20}([0-9]+(?:\.[0-9]+)?)\s*%",
        r"提点比率最高限价[：:\s]*[^0-9\n]{0,20}([0-9]+(?:\.[0-9]+)?)\s*%",
    ],
    "bid_deadline": [
        r"提交投标文件截止时间[：:\s]+([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日[^。\n]*)",
        r"投标截止时间[：:\s]+([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日[^。\n]*)",
        r"开标时间[：:\s]+([0-9]{4}年[0-9]{1,2}月[0-9]{1,2}日[^。\n]*)",
    ],
}


def normalize_text(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def read_file_index(path: Path) -> list[dict]:
    with path.open("r", encoding="utf-8-sig", newline="") as f:
        return list(csv.DictReader(f))


def docx_text(path: Path) -> str:
    doc = Document(path)
    parts = []
    parts.extend(p.text for p in doc.paragraphs if normalize_text(p.text))
    for table in doc.tables:
        for row in table.rows:
            cells = [normalize_text(cell.text) for cell in row.cells if normalize_text(cell.text)]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def pdf_text(path: Path) -> str:
    try:
        import pdfplumber
    except Exception:
        return ""
    texts = []
    try:
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages[:20]:
                texts.append(page.extract_text() or "")
    except Exception:
        return ""
    return "\n".join(texts)


def read_text(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        return docx_text(path)
    if suffix == ".pdf":
        return pdf_text(path)
    return ""


def clean_amount_to_yuan(value: str) -> float | None:
    text = value.replace(",", "").replace("，", "")
    match = re.search(r"[0-9]+(?:\.[0-9]+)?", text)
    if not match:
        return None
    amount = float(match.group(0))
    return round(amount * 10000 if "万" in value else amount, 2)


def clean_rate(value: str) -> float | None:
    match = re.search(r"[0-9]+(?:\.[0-9]+)?", value)
    if not match:
        return None
    return float(match.group(0))


def extract_field(text: str, field: str) -> str:
    for pattern in FIELD_PATTERNS[field]:
        match = re.search(pattern, text)
        if match:
            return normalize_text(match.group(1))
    return ""


def contains_any(text: str, words: list[str]) -> bool:
    return any(word in text for word in words)


def extract_sections(text: str) -> dict[str, bool]:
    return {
        "has_scoring_method": contains_any(text, ["评分办法", "综合评分", "评标办法", "评分标准"]),
        "has_qualification": contains_any(text, ["资格要求", "申请人的资格要求", "供应商资格", "投标人资格"]),
        "has_contract_terms": contains_any(text, ["合同条款", "合同格式", "合同主要条款"]),
        "has_bid_format": contains_any(text, ["投标文件格式", "响应文件格式", "资格声明函"]),
        "has_procurement_need": contains_any(text, ["采购需求", "技术要求", "服务要求", "项目需求"]),
    }


def parse_tender_file(path: Path, source_row: dict, base_dir: Path) -> dict:
    text = read_text(path)
    budget_raw = extract_field(text, "budget_amount")
    ceiling_raw = extract_field(text, "ceiling_price")
    ceiling_rate_raw = extract_field(text, "ceiling_rate")
    sections = extract_sections(text)
    return {
        "file_path": str(path.relative_to(base_dir)),
        "file_name": path.name,
        "extension": path.suffix.lower(),
        "text_length": len(text),
        "project_name": extract_field(text, "project_name"),
        "project_no": extract_field(text, "project_no"),
        "purchaser": extract_field(text, "purchaser") or source_row.get("purchaser", ""),
        "budget_amount_raw": budget_raw,
        "budget_amount_yuan": clean_amount_to_yuan(budget_raw) if budget_raw else "",
        "ceiling_price_raw": ceiling_raw,
        "ceiling_price_yuan": clean_amount_to_yuan(ceiling_raw) if ceiling_raw else "",
        "ceiling_rate_raw": ceiling_rate_raw,
        "ceiling_rate_pct": clean_rate(ceiling_rate_raw) if ceiling_rate_raw else "",
        "bid_deadline": extract_field(text, "bid_deadline"),
        **sections,
        "notice_title": source_row.get("notice_title", ""),
        "notice_url": source_row.get("notice_url", ""),
        "attachment_name": source_row.get("attachment_name", ""),
        "text_excerpt": normalize_text(text[:800]),
    }


def parse_files(file_index_path: Path, attachment_root: Path, extracted_root: Path, output_dir: Path) -> list[dict]:
    rows = read_file_index(file_index_path)
    parsed = []
    for row in rows:
        if str(row.get("is_document", "")).lower() != "true":
            continue
        rel_path = row["file_path"]
        root = extracted_root if row.get("extracted_from") else attachment_root
        path = root / rel_path
        if not path.exists():
            continue
        if path.suffix.lower() not in {".docx", ".pdf"}:
            continue
        parsed.append(parse_tender_file(path, row, root))

    output_dir.mkdir(parents=True, exist_ok=True)
    write_jsonl(output_dir / "parsed_tender_files.jsonl", parsed)
    write_csv(output_dir / "parsed_tender_files.csv", parsed)
    summary = {
        "file_count": len(parsed),
        "with_project_name": sum(1 for item in parsed if item["project_name"]),
        "with_project_no": sum(1 for item in parsed if item["project_no"]),
        "with_budget": sum(1 for item in parsed if item["budget_amount_raw"]),
        "with_ceiling": sum(1 for item in parsed if item["ceiling_price_raw"]),
        "with_ceiling_rate": sum(1 for item in parsed if item["ceiling_rate_raw"]),
        "with_deadline": sum(1 for item in parsed if item["bid_deadline"]),
    }
    (output_dir / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    return parsed


def write_jsonl(path: Path, rows: list[dict]) -> None:
    with path.open("w", encoding="utf-8") as f:
        for row in rows:
            f.write(json.dumps(row, ensure_ascii=False) + "\n")


def write_csv(path: Path, rows: list[dict]) -> None:
    fieldnames = list(rows[0].keys()) if rows else [
        "file_path",
        "file_name",
        "extension",
        "text_length",
        "project_name",
        "project_no",
        "purchaser",
        "budget_amount_raw",
        "budget_amount_yuan",
        "ceiling_price_raw",
        "ceiling_price_yuan",
        "ceiling_rate_raw",
        "ceiling_rate_pct",
        "bid_deadline",
        "has_scoring_method",
        "has_qualification",
        "has_contract_terms",
        "has_bid_format",
        "has_procurement_need",
        "notice_title",
        "notice_url",
        "attachment_name",
        "text_excerpt",
    ]
    with path.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows)


def main() -> None:
    parser = argparse.ArgumentParser(description="Parse public tender files into structured training samples.")
    parser.add_argument("--file-index", required=True, help="file_index.csv generated by index_public_attachments.")
    parser.add_argument("--attachment-root", required=True, help="Downloaded attachment root.")
    parser.add_argument("--extracted-root", required=True, help="Extracted archive root.")
    parser.add_argument("--output", required=True, help="Output directory.")
    args = parser.parse_args()
    parsed = parse_files(
        file_index_path=Path(args.file_index),
        attachment_root=Path(args.attachment_root),
        extracted_root=Path(args.extracted_root),
        output_dir=Path(args.output),
    )
    print(f"Parsed {len(parsed)} tender files into {args.output}")


if __name__ == "__main__":
    main()
