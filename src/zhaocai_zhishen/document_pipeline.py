from __future__ import annotations

import collections
import csv
import os
import hashlib
import io
import json
import re
from dataclasses import asdict, dataclass, field
from pathlib import Path
from typing import Iterable, Protocol

from docx import Document
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".docx", ".pdf"}
MIN_TEXT_CHARS_PER_PAGE = 40
CACHE_VERSION = "v4"
OCR_ENGINE_CHOICES = {"auto", "paddle-gpu", "rapidocr"}
_HEADER_FOOTER_WINDOW = 4
_MIN_REPEAT_LINE_LENGTH = 6


def normalize_text(value: object) -> str:
    text = str(value or "").replace("\u3000", " ").replace("\x00", "")
    lines = [re.sub(r"[ \t]+", " ", line).strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line)


def _line_key(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


def _is_numeric_or_page_marker(line: str) -> bool:
    return bool(re.fullmatch(r"[-—]?(?:第\s*)?\d{1,4}(?:\s*(?:页|/|of)\s*\d{0,4})?", line, flags=re.IGNORECASE))


def clean_extracted_pages(pages: list["ExtractedPage"]) -> dict[str, int]:
    """Remove OCR boilerplate from analysis text while retaining raw page text for evidence replay."""
    normalized_pages: list[list[str]] = []
    page_line_counts: collections.Counter[str] = collections.Counter()
    shell_line_counts: collections.Counter[str] = collections.Counter()
    page_signatures: collections.Counter[str] = collections.Counter()

    for page in pages:
        lines = [_line_key(line) for line in normalize_text(page.text).splitlines() if _line_key(line)]
        normalized_pages.append(lines)
        for line in set(lines):
            page_line_counts[line] += 1
        for index, line in enumerate(lines):
            if index < _HEADER_FOOTER_WINDOW or index >= max(0, len(lines) - _HEADER_FOOTER_WINDOW):
                shell_line_counts[line] += 1
        compact = "\n".join(lines)
        if len(compact) >= 80:
            page_signatures[hashlib.sha256(compact.encode("utf-8")).hexdigest()] += 1

    page_count = len(pages)
    repeated_shell_lines = {
        line
        for line, count in page_line_counts.items()
        if count >= max(3, round(page_count * 0.02))
        and shell_line_counts[line] / max(count, 1) >= 0.6
        and (len(line) <= 120 or _is_numeric_or_page_marker(line))
    }
    seen_shell_lines: set[str] = set()
    seen_page_signatures: set[str] = set()
    removed_lines = 0
    duplicate_pages = 0

    for page, lines in zip(pages, normalized_pages):
        compact = "\n".join(lines)
        signature = hashlib.sha256(compact.encode("utf-8")).hexdigest() if len(compact) >= 80 else ""
        if signature and page_signatures[signature] > 1 and signature in seen_page_signatures:
            page.cleaned_text = ""
            page.removed_redundant_line_count = len(lines)
            removed_lines += len(lines)
            duplicate_pages += 1
            continue
        if signature:
            seen_page_signatures.add(signature)

        kept: list[str] = []
        seen_in_page: set[str] = set()
        for line in lines:
            if line in repeated_shell_lines:
                if line in seen_shell_lines:
                    removed_lines += 1
                    continue
                seen_shell_lines.add(line)
            if len(line) >= _MIN_REPEAT_LINE_LENGTH and line in seen_in_page and not _is_numeric_or_page_marker(line):
                removed_lines += 1
                continue
            kept.append(line)
            seen_in_page.add(line)
        page.cleaned_text = "\n".join(kept)
        page.removed_redundant_line_count = len(lines) - len(kept)

    return {
        "raw_line_count": sum(len(lines) for lines in normalized_pages),
        "cleaned_line_count": sum(len(page.cleaned_text.splitlines()) for page in pages),
        "removed_redundant_line_count": removed_lines,
        "duplicate_page_count": duplicate_pages,
        "repeated_shell_line_count": len(repeated_shell_lines),
    }


def stable_document_id(path: Path, sample_id: str = "") -> str:
    if sample_id:
        return sample_id
    stat = path.stat()
    value = f"{path.resolve()}|{stat.st_size}|{stat.st_mtime_ns}"
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:24]


@dataclass
class SourceSample:
    document_id: str
    project_id: str
    bidder_name: str
    path: Path
    source_inner_path: str = ""


@dataclass
class ExtractedPage:
    page_number: int
    text: str
    char_count: int
    extraction_method: str
    cleaned_text: str = ""
    removed_redundant_line_count: int = 0


@dataclass
class ExtractedDocument:
    document_id: str
    project_id: str
    bidder_name: str
    file_path: str
    file_type: str
    file_size: int
    page_count: int
    text: str
    text_char_count: int
    is_scanned: bool
    needs_ocr: bool
    parse_status: str
    extraction_method: str
    source_inner_path: str = ""
    warnings: list[str] = field(default_factory=list)
    raw_text_char_count: int = 0
    removed_redundant_line_count: int = 0
    duplicate_page_count: int = 0


class OcrEngine(Protocol):
    def extract_pdf(self, path: Path, max_pages: int | None = None) -> list[ExtractedPage]: ...


class RapidOcrEngine:
    """优先复用 PDF 内嵌的整页图片，避免额外安装 PDF 渲染器。"""

    def __init__(self) -> None:
        try:
            from PIL import Image
            from rapidocr_onnxruntime import RapidOCR
        except ImportError as exc:
            raise RuntimeError("RapidOCR 未安装") from exc
        self._ocr = RapidOCR(
            use_cls=False,
            rec_batch_num=32,
            intra_op_num_threads=int(os.environ.get("BID_AUDIT_OCR_THREADS", "8")),
            inter_op_num_threads=1,
        )
        self._max_image_width = 1800
        self._resampling = Image.Resampling.LANCZOS

    def extract_pdf(self, path: Path, max_pages: int | None = None) -> list[ExtractedPage]:
        reader = PdfReader(path)
        source_pages = list(reader.pages)
        if max_pages is not None:
            source_pages = source_pages[:max_pages]
        pages: list[ExtractedPage] = []
        for index, page in enumerate(source_pages):
            images = list(page.images)
            if not images:
                pages.append(ExtractedPage(index + 1, "", 0, "rapidocr"))
                continue
            image = max(images, key=lambda item: item.image.width * item.image.height).image.convert("RGB")
            if image.width > self._max_image_width:
                height = round(image.height * self._max_image_width / image.width)
                image = image.resize((self._max_image_width, height), self._resampling)
            buffer = io.BytesIO()
            image.save(buffer, format="JPEG", quality=92)
            result, _ = self._ocr(buffer.getvalue(), use_cls=False)
            text = normalize_text("\n".join(str(item[1]) for item in (result or []) if len(item) > 1))
            pages.append(ExtractedPage(index + 1, text, len(text), "rapidocr"))
        return pages


class PaddleOcrEngine:
    """使用本机 NVIDIA GPU 的 PaddleOCR 适配器。"""

    def __init__(self) -> None:
        os.environ.setdefault("PADDLE_PDX_MODEL_SOURCE", "modelscope")
        os.environ.setdefault("PADDLE_PDX_DISABLE_MODEL_SOURCE_CHECK", "True")
        try:
            import fitz
            import numpy as np
            import paddle
            from paddleocr import PaddleOCR
        except ImportError as exc:
            raise RuntimeError("PaddleOCR GPU 依赖未安装") from exc
        if not paddle.is_compiled_with_cuda() or paddle.device.cuda.device_count() < 1:
            raise RuntimeError("当前 PaddlePaddle 未启用 CUDA 或未检测到 NVIDIA GPU")
        paddle.set_device("gpu:0")
        self._fitz = fitz
        self._np = np
        self._paddle = paddle
        self._ocr = PaddleOCR(
            device="gpu:0",
            text_detection_model_name="PP-OCRv5_mobile_det",
            text_recognition_model_name="PP-OCRv5_mobile_rec",
            text_recognition_batch_size=16,
            use_doc_orientation_classify=False,
            use_doc_unwarping=False,
            use_textline_orientation=False,
        )

    def extract_pdf(self, path: Path, max_pages: int | None = None) -> list[ExtractedPage]:
        pages: list[ExtractedPage] = []
        pdf = self._fitz.open(path)
        try:
            source_pages = list(pdf)
            if max_pages is not None:
                source_pages = source_pages[:max_pages]
            for index, page in enumerate(source_pages):
                pixmap = page.get_pixmap(matrix=self._fitz.Matrix(2, 2), alpha=False)
                image = self._np.frombuffer(pixmap.samples, dtype=self._np.uint8)
                image = image.reshape(pixmap.height, pixmap.width, pixmap.n)
                if pixmap.n == 4:
                    image = image[:, :, :3]
                image = image[:, :, ::-1].copy()
                result = self._ocr.predict(image)
                parts: list[str] = []
                for item in result:
                    payload = getattr(item, "json", None)
                    payload = payload if isinstance(payload, dict) else {}
                    data = payload.get("res", payload)
                    parts.extend(str(value) for value in data.get("rec_texts", []) if value)
                text = normalize_text("\n".join(parts))
                pages.append(ExtractedPage(index + 1, text, len(text), "paddleocr"))
                del result, image, pixmap
                if (index + 1) % 10 == 0:
                    self._paddle.device.cuda.empty_cache()
        finally:
            pdf.close()
            self._paddle.device.cuda.empty_cache()
        return pages


def create_ocr_engine(engine_name: str = "auto") -> OcrEngine:
    if engine_name not in OCR_ENGINE_CHOICES:
        raise ValueError(f"不支持的 OCR 引擎: {engine_name}")
    errors: list[str] = []
    engine_types = {
        "paddle-gpu": (PaddleOcrEngine,),
        "rapidocr": (RapidOcrEngine,),
        "auto": (PaddleOcrEngine, RapidOcrEngine),
    }[engine_name]
    for engine_type in engine_types:
        try:
            return engine_type()
        except RuntimeError as exc:
            errors.append(str(exc))
    raise RuntimeError("；".join(errors))


def extract_docx(path: Path) -> tuple[list[ExtractedPage], list[str]]:
    doc = Document(path)
    paragraphs = [normalize_text(paragraph.text) for paragraph in doc.paragraphs]
    table_rows: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            values = [normalize_text(cell.text) for cell in row.cells]
            table_rows.append(" | ".join(value for value in values if value))
    text = normalize_text("\n".join(paragraphs + table_rows))
    return [ExtractedPage(1, text, len(text), "python-docx")], []


def extract_pdf(path: Path, ocr_engine: OcrEngine | None = None, max_pages: int | None = None) -> tuple[list[ExtractedPage], list[str]]:
    reader = PdfReader(path)
    source_pages = list(reader.pages)
    if max_pages is not None:
        source_pages = source_pages[:max_pages]
    pages: list[ExtractedPage] = []
    warnings: list[str] = []
    for index, page in enumerate(source_pages):
        try:
            text = normalize_text(page.extract_text() or "")
        except Exception as exc:
            text = ""
            warnings.append(f"第 {index + 1} 页文本提取失败: {type(exc).__name__}")
        pages.append(ExtractedPage(index + 1, text, len(text), "pypdf"))

    sparse_pages = sum(page.char_count < MIN_TEXT_CHARS_PER_PAGE for page in pages)
    is_scanned = bool(pages) and sparse_pages / len(pages) >= 0.8
    if is_scanned and ocr_engine is not None:
        ocr_pages = ocr_engine.extract_pdf(path, max_pages)
        if sum(page.char_count for page in ocr_pages) > sum(page.char_count for page in pages):
            pages = ocr_pages
        else:
            warnings.append("OCR 文本量未超过原生 PDF 提取结果")
    return pages, warnings


def load_samples(dataset_dir: Path) -> list[SourceSample]:
    manifest = dataset_dir / "samples.csv"
    if not manifest.exists():
        raise FileNotFoundError(f"未找到样本清单: {manifest}")
    samples: list[SourceSample] = []
    with manifest.open("r", encoding="utf-8-sig", newline="") as handle:
        for row in csv.DictReader(handle):
            raw_path = (row.get("bid_file") or "").replace("\\", "/")
            path = dataset_dir / Path(raw_path)
            if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue
            samples.append(SourceSample(
                document_id=stable_document_id(path, row.get("sample_id") or ""),
                project_id=row.get("project_name") or row.get("archive_id") or path.parent.parent.name,
                bidder_name=row.get("bidder_candidate") or path.stem,
                path=path,
                source_inner_path=row.get("source_inner_path") or "",
            ))
    return samples


def discover_samples(input_dir: Path) -> list[SourceSample]:
    samples: list[SourceSample] = []
    for path in sorted(input_dir.rglob("*")):
        if not path.is_file() or path.suffix.lower() not in SUPPORTED_EXTENSIONS or path.name.startswith("~$"):
            continue
        samples.append(SourceSample(
            document_id=stable_document_id(path),
            project_id=path.parent.parent.name if len(path.parents) > 1 else input_dir.name,
            bidder_name=path.stem,
            path=path,
        ))
    return samples


def extract_sample(sample: SourceSample, root: Path, ocr_engine: OcrEngine | None = None, max_pages: int | None = None) -> tuple[ExtractedDocument, list[ExtractedPage]]:
    if not sample.path.exists():
        raise FileNotFoundError(sample.path)
    extension = sample.path.suffix.lower()
    if extension == ".docx":
        pages, warnings = extract_docx(sample.path)
    elif extension == ".pdf":
        pages, warnings = extract_pdf(sample.path, ocr_engine, max_pages)
    else:
        raise ValueError(f"不支持的文件类型: {extension}")

    cleaning = clean_extracted_pages(pages)
    raw_text = normalize_text("\n\n".join(page.text for page in pages))
    text = normalize_text("\n\n".join(page.cleaned_text for page in pages))
    sparse_pages = sum(page.char_count < MIN_TEXT_CHARS_PER_PAGE for page in pages)
    used_ocr = any(page.extraction_method in {"rapidocr", "paddleocr"} for page in pages)
    is_scanned = extension == ".pdf" and (used_ocr or (bool(pages) and sparse_pages / len(pages) >= 0.8))
    needs_ocr = is_scanned and not used_ocr
    document = ExtractedDocument(
        document_id=sample.document_id,
        project_id=sample.project_id,
        bidder_name=sample.bidder_name,
        file_path=sample.path.relative_to(root).as_posix(),
        file_type=extension.lstrip("."),
        file_size=sample.path.stat().st_size,
        page_count=len(pages),
        text=text,
        text_char_count=len(text),
        is_scanned=is_scanned,
        needs_ocr=needs_ocr,
        parse_status="needs_ocr" if needs_ocr else "success",
        extraction_method="+".join(sorted({page.extraction_method for page in pages})),
        source_inner_path=sample.source_inner_path,
        warnings=warnings,
        raw_text_char_count=len(raw_text),
        removed_redundant_line_count=cleaning["removed_redundant_line_count"],
        duplicate_page_count=cleaning["duplicate_page_count"],
    )
    return document, pages


_WORKER_OCR_ENGINE: OcrEngine | None = None


def initialize_ocr_worker(use_ocr: bool, threads: int, ocr_engine_name: str) -> None:
    global _WORKER_OCR_ENGINE
    os.environ["BID_AUDIT_OCR_THREADS"] = str(threads)
    _WORKER_OCR_ENGINE = create_ocr_engine(ocr_engine_name) if use_ocr else None


def extract_sample_worker(payload: tuple[int, SourceSample, Path, int | None]) -> tuple[int, dict, list[dict]]:
    index, sample, root, max_pages = payload
    document, pages = extract_sample(sample, root, _WORKER_OCR_ENGINE, max_pages)
    return index, asdict(document), [asdict(page) for page in pages]


def write_jsonl(path: Path, rows: Iterable[dict]) -> None:
    with path.open("w", encoding="utf-8", newline="\n") as handle:
        for row in rows:
            handle.write(json.dumps(row, ensure_ascii=False) + "\n")


def write_cache(path: Path, payload: dict) -> None:
    temporary = path.with_suffix(path.suffix + ".tmp")
    temporary.write_text(json.dumps(payload, ensure_ascii=False), encoding="utf-8")
    temporary.replace(path)


def cache_path_for(
    sample: SourceSample,
    cache_dir: Path,
    ocr_mode: str,
    ocr_engine_name: str,
    max_pages: int | None,
) -> Path:
    stat = sample.path.stat()
    key = (
        f"{sample.document_id}.{CACHE_VERSION}.{stat.st_size}.{stat.st_mtime_ns}."
        f"{ocr_mode}.{ocr_engine_name}.{max_pages or 'all'}.json"
    )
    return cache_dir / key


def prepare_dataset(
    input_dir: Path,
    output_dir: Path,
    *,
    ocr_mode: str = "auto",
    ocr_engine_name: str = "auto",
    max_documents: int | None = None,
    max_pages_per_document: int | None = None,
    reuse_cache: bool = True,
    workers: int = 1,
) -> dict[str, int | str]:
    from concurrent.futures import ProcessPoolExecutor, as_completed

    input_dir = input_dir.resolve()
    output_dir.mkdir(parents=True, exist_ok=True)
    cache_dir = output_dir / "cache"
    cache_dir.mkdir(exist_ok=True)
    samples = load_samples(input_dir) if (input_dir / "samples.csv").exists() else discover_samples(input_dir)
    if max_documents is not None:
        samples = samples[:max_documents]
    workers = max(1, workers)

    use_ocr = ocr_mode in {"auto", "on"}
    ocr_engine: OcrEngine | None = None
    ocr_warning = ""
    if use_ocr:
        try:
            ocr_engine = create_ocr_engine(ocr_engine_name)
        except RuntimeError as exc:
            if ocr_mode == "on":
                raise
            use_ocr = False
            ocr_warning = str(exc)

    selected_ocr_engine = "off"
    if isinstance(ocr_engine, PaddleOcrEngine):
        selected_ocr_engine = "paddle-gpu"
        if workers > 1:
            workers = 1
            ocr_warning = "GPU OCR 使用单进程，已将 workers 自动调整为 1"
    elif isinstance(ocr_engine, RapidOcrEngine):
        selected_ocr_engine = "rapidocr"

    results: dict[int, tuple[dict, list[dict]]] = {}
    failures: list[dict] = []
    pending: list[tuple[int, SourceSample, Path]] = []
    cached_count = 0
    completed_count = 0

    for index, sample in enumerate(samples):
        cache_path = cache_path_for(
            sample,
            cache_dir,
            ocr_mode,
            selected_ocr_engine,
            max_pages_per_document,
        )
        if reuse_cache and cache_path.exists():
            try:
                cached = json.loads(cache_path.read_text(encoding="utf-8"))
                results[index] = (cached["document"], cached["pages"])
                cached_count += 1
                completed_count += 1
                print(f"[{completed_count}/{len(samples)}] 缓存", flush=True)
                continue
            except (OSError, KeyError, json.JSONDecodeError):
                cache_path.unlink(missing_ok=True)
        pending.append((index, sample, cache_path))

    if workers > 1 and pending:
        threads = max(1, min(8, 16 // workers))
        payloads = [(index, sample, input_dir, max_pages_per_document) for index, sample, _ in pending]
        cache_paths = {index: cache_path for index, _, cache_path in pending}
        sample_paths = {index: sample.path for index, sample, _ in pending}
        with ProcessPoolExecutor(
            max_workers=workers,
            initializer=initialize_ocr_worker,
            initargs=(use_ocr, threads, selected_ocr_engine),
        ) as executor:
            future_map = {executor.submit(extract_sample_worker, payload): payload[0] for payload in payloads}
            for future in as_completed(future_map):
                index = future_map[future]
                completed_count += 1
                try:
                    _, document_row, pages = future.result()
                    results[index] = (document_row, pages)
                    write_cache(cache_paths[index], {"document": document_row, "pages": pages})
                    print(f"[{completed_count}/{len(samples)}] {document_row['parse_status']}", flush=True)
                except Exception as exc:
                    failures.append({
                        "document_id": samples[index].document_id,
                        "file_path": str(sample_paths[index]),
                        "error_type": type(exc).__name__,
                        "error": str(exc),
                    })
                    print(f"[{completed_count}/{len(samples)}] 失败: {type(exc).__name__}", flush=True)
    else:
        for index, sample, cache_path in pending:
            completed_count += 1
            try:
                document, extracted_pages = extract_sample(sample, input_dir, ocr_engine if use_ocr else None, max_pages_per_document)
                document_row = asdict(document)
                pages = [asdict(page) for page in extracted_pages]
                results[index] = (document_row, pages)
                write_cache(cache_path, {"document": document_row, "pages": pages})
                print(f"[{completed_count}/{len(samples)}] {document.parse_status}", flush=True)
            except Exception as exc:
                failures.append({
                    "document_id": sample.document_id,
                    "file_path": str(sample.path),
                    "error_type": type(exc).__name__,
                    "error": str(exc),
                })
                print(f"[{completed_count}/{len(samples)}] 失败: {type(exc).__name__}", flush=True)

    documents: list[dict] = []
    page_rows: list[dict] = []
    for index in sorted(results):
        document_row, pages = results[index]
        documents.append(document_row)
        page_rows.extend({"document_id": document_row["document_id"], **page} for page in pages)

    write_jsonl(output_dir / "documents.jsonl", documents)
    write_jsonl(output_dir / "pages.jsonl", page_rows)
    write_jsonl(output_dir / "unsupervised_samples.jsonl", documents)
    with (output_dir / "parse_failures.csv").open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=["document_id", "file_path", "error_type", "error"])
        writer.writeheader()
        writer.writerows(failures)

    summary = {
        "input_dir": str(input_dir),
        "output_dir": str(output_dir.resolve()),
        "sample_count": len(samples),
        "success_count": len(documents),
        "cached_count": cached_count,
        "failure_count": len(failures),
        "needs_ocr_count": sum(bool(row["needs_ocr"]) for row in documents),
        "page_count": sum(int(row["page_count"]) for row in documents),
        "text_char_count": sum(int(row["text_char_count"]) for row in documents),
        "raw_text_char_count": sum(int(row.get("raw_text_char_count", row["text_char_count"])) for row in documents),
        "removed_redundant_line_count": sum(int(row.get("removed_redundant_line_count", 0)) for row in documents),
        "duplicate_page_count": sum(int(row.get("duplicate_page_count", 0)) for row in documents),
        "source_size_bytes": sum(int(row["file_size"]) for row in documents),
        "workers": workers,
        "ocr_engine": selected_ocr_engine,
        "ocr_warning": ocr_warning,
    }
    (output_dir / "summary.json").write_text(
        json.dumps(summary, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return summary
