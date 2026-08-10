from __future__ import annotations

import re
from pathlib import Path

from docx import Document

from .models import BidDocument


BID_FOLDER_NAMES = {"投标文件", "投标书", "投标资料"}


def normalize_spaces(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def clean_number(value: object) -> float | None:
    text = str(value or "")
    text = text.replace(",", "").replace("，", "").replace("¥", "").replace("￥", "")
    match = re.search(r"\d+(?:\.\d+)?", text)
    if not match:
        return None
    try:
        number = float(match.group(0))
    except ValueError:
        return None
    return number if number > 0 else None


def row_values(row) -> list[str]:
    return [normalize_spaces(cell.text.replace("\n", " ")) for cell in row.cells]


def after_prefix(lines: list[str], prefix: str) -> str:
    for line in lines[:80]:
        if prefix in line:
            return line.split(prefix, 1)[1].lstrip(" ：:").strip()
    return ""


def extract_project(lines: list[str]) -> str:
    for line in lines[:24]:
        if "项目" in line and "投标文件" not in line and "招标编号" not in line:
            return line
    return ""


def extract_price_table(doc) -> dict[str, float | None]:
    result: dict[str, float | None] = {"total": None, "device": None, "service": None, "install": None}
    for table in doc.tables[:4]:
        for row in table.rows[:24]:
            values = row_values(row)
            joined = " ".join(values)
            nums = [clean_number(v) for v in values]
            nums = [n for n in nums if n and n > 1000]
            if "投标总价" in joined and nums:
                result["total"] = max(nums)
            if "设备价格" in joined and nums:
                result["device"] = nums[-1]
            if "售后服务" in joined and nums:
                result["service"] = nums[-1]
            if "安装调试" in joined and nums:
                result["install"] = nums[-1]
        if result["total"]:
            return result
    return result


def extract_company_table(doc) -> dict[str, str]:
    info = {"contact": "", "phone": "", "address": "", "capital": "", "founded": "", "staff": ""}
    for table in doc.tables:
        flat = " ".join(cell.text for row in table.rows for cell in row.cells)
        if "投标人名称" not in flat or "联系方式" not in flat:
            continue
        for row in table.rows:
            values = row_values(row)
            if values and "注册地址" in values[0] and len(values) > 1:
                info["address"] = values[1]
            if values and "注册资金" in values[0] and len(values) > 1:
                info["capital"] = values[1]
            if "成立时间" in values:
                idx = values.index("成立时间")
                if idx + 1 < len(values):
                    info["founded"] = values[idx + 1]
            if "员工总数" in values:
                idx = values.index("员工总数")
                if idx + 1 < len(values):
                    info["staff"] = values[idx + 1]
            if "联系方式" in values and "联系人" in values and len(values) >= 5:
                info["contact"] = values[2]
                info["phone"] = values[4]
        return info
    return info


def extract_references(doc) -> list[str]:
    refs: list[str] = []
    for table in doc.tables:
        for row in table.rows:
            values = row_values(row)
            joined = " ".join(values)
            if "项目名称" in joined and len(joined) > 8:
                refs.append(joined[:160])
            if len(refs) >= 4:
                return refs
    return refs


def extract_warranty(doc) -> str:
    for table in doc.tables[:5]:
        header = " ".join(row_values(table.rows[0])) if table.rows else ""
        if "质保期" not in header:
            continue
        values = []
        for row in table.rows[1:]:
            cells = row_values(row)
            if cells:
                values.append(cells[-1])
        values = [v for v in values if v]
        return "；".join(sorted(set(values)))[:80]
    return ""


def parse_docx(path: Path, data_dir: Path, dataset_project: str) -> BidDocument:
    doc = Document(path)
    lines = [normalize_spaces(p.text) for p in doc.paragraphs if normalize_spaces(p.text)]
    table_lines = [
        normalize_spaces(cell.text)
        for table in doc.tables
        for row in table.rows
        for cell in row.cells
        if normalize_spaces(cell.text)
    ]
    full_text = "\n".join(lines + table_lines)
    props = doc.core_properties
    price = extract_price_table(doc)
    company = extract_company_table(doc)
    bidder = after_prefix(lines, "投标人") or path.parent.name
    bidder = bidder.replace("（盖单位章）", "").replace("(盖单位章)", "").strip()
    if len(bidder) < 6:
        bidder = path.parent.name
    dates = re.findall(r"20\d{2}年\d{1,2}月\d{1,2}日", full_text[:1600])
    return BidDocument(
        dataset_project=dataset_project,
        file=str(path.relative_to(data_dir)),
        folder=path.parent.name,
        bidder=bidder,
        project=extract_project(lines) or "未识别项目",
        tender_no=after_prefix(lines, "招标编号"),
        bid_date=dates[0] if dates else "",
        price=price["total"],
        device_price=price["device"],
        service_price=price["service"],
        install_price=price["install"],
        contact=company["contact"],
        phone=company["phone"],
        address=company["address"],
        capital=company["capital"],
        founded=company["founded"],
        staff=company["staff"],
        warranty=extract_warranty(doc),
        negative_deviations=full_text.count("负偏离"),
        neutral_deviations=full_text.count("无偏离"),
        table_count=len(doc.tables),
        paragraph_count=len(lines),
        author=str(props.author or ""),
        last_modified_by=str(props.last_modified_by or ""),
        created=str(props.created or ""),
        modified=str(props.modified or ""),
        revision=str(props.revision or ""),
        references=extract_references(doc),
    )


def discover_project_sources(data_dir: Path) -> list[tuple[str, Path]]:
    """Return (project_name, search_root) pairs.

    Supported layouts:
    1. Single-project legacy layout:
       data_dir/
         bidder_a/*.docx
         bidder_b/*.docx

    2. Multi-project layout:
       data_dir/
         project_a/投标文件/*.docx
         project_b/投标文件/*.docx
    """
    child_dirs = [p for p in data_dir.iterdir() if p.is_dir()]
    project_sources: list[tuple[str, Path]] = []
    for child in child_dirs:
        bid_subdirs = [p for p in child.iterdir() if p.is_dir() and p.name in BID_FOLDER_NAMES]
        for bid_subdir in bid_subdirs:
            if any(p.suffix.lower() == ".docx" and not p.name.startswith("~$") for p in bid_subdir.rglob("*.docx")):
                project_sources.append((child.name, bid_subdir))
                break

    if project_sources:
        return project_sources

    return [(data_dir.name, data_dir)]


def parse_directory(data_dir: Path) -> list[BidDocument]:
    documents: list[BidDocument] = []
    for dataset_project, search_root in discover_project_sources(data_dir):
        files = [p for p in search_root.rglob("*.docx") if not p.name.startswith("~$")]
        for path in files:
            try:
                documents.append(parse_docx(path, data_dir, dataset_project))
            except Exception as exc:
                print(f"解析失败: {path} - {exc}")
    return documents
