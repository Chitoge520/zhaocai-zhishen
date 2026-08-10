from __future__ import annotations

import html
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Inches, Pt

from .analysis_results import load_unsupervised_results
from .evidence_graph import build_evidence_graph


def _read_json(path: Path) -> dict:
    if not path.exists():
        return {}
    try:
        value = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, UnicodeError, json.JSONDecodeError):
        return {}
    return value if isinstance(value, dict) else {}


def build_report_payload(
    analysis_dir: Path,
    processed_dir: Path,
    training_dir: Path,
    models_dir: Path,
    *,
    title: str = "招采智审异常线索复核报告",
    llm: dict | None = None,
) -> dict:
    analysis = load_unsupervised_results(
        analysis_dir,
        model_path=models_dir / "bid_anomaly_model.json",
    )
    processed = _read_json(processed_dir / "summary.json")
    training = _read_json(training_dir / "summary.json")
    model = _read_json(models_dir / "bid_anomaly_model.json")
    model_summary = _read_json(models_dir / "training_summary.json")
    graph = build_evidence_graph(analysis.get("entities", []), analysis.get("pairs", []), analysis.get("anomalies", []))
    summary = analysis.get("summary", {})
    return {
        "title": title,
        "generated_at": datetime.now().astimezone().isoformat(timespec="seconds"),
        "analysis": analysis,
        "summary": summary,
        "processed": processed,
        "training": training,
        "model": model,
        "model_summary": model_summary,
        "graph": graph,
        "llm": llm or {"status": "skipped", "findings": [], "validated_finding_count": 0},
        "boundary": "本报告输出异常线索和待复核证据，不直接认定串标、围标或其他违规行为。",
    }


def _esc(value: object) -> str:
    return html.escape(str(value if value is not None else ""))


def _pct(value: object) -> str:
    try:
        return f"{float(value or 0) * 100:.2f}%"
    except (TypeError, ValueError):
        return "0.00%"


def _metric_rows(payload: dict) -> list[tuple[str, object]]:
    summary = payload.get("summary", {})
    processed = payload.get("processed", {})
    training = payload.get("training", {})
    graph = payload.get("graph", {})
    return [
        ("历史/本次项目数", summary.get("project_count", graph.get("project_count", 0))),
        ("投标文件数", processed.get("success_count", summary.get("document_count", 0))),
        ("解析页数", processed.get("page_count", 0)),
        ("原始压缩包数", training.get("archive_count", 0)),
        ("项目内比较数", summary.get("pair_count", 0)),
        ("待复核线索数", summary.get("anomaly_count", len(payload.get("analysis", {}).get("anomalies", [])))),
        ("重复片段证据数", summary.get("repeated_segment_evidence_count", 0)),
        ("公共模板片段排除数", summary.get("common_template_segment_count", 0)),
    ]


def render_html_report(payload: dict) -> bytes:
    summary = payload.get("summary", {})
    model_summary = payload.get("model_summary", {})
    graph = payload.get("graph", {})
    llm = payload.get("llm", {})
    findings = payload.get("analysis", {}).get("anomalies", [])
    metrics = "".join(
        f'<div class="metric"><span>{_esc(label)}</span><strong>{_esc(value)}</strong></div>'
        for label, value in _metric_rows(payload)
    )
    finding_sections = []
    for index, finding in enumerate(findings, start=1):
        evidence = "".join(f"<li>{_esc(item)}</li>" for item in finding.get("evidence", []))
        finding_sections.append(
            f"""
            <section class="finding">
              <h2>{index}. {_esc(finding.get('bidder_a'))} 对比 {_esc(finding.get('bidder_b'))}</h2>
               <p><b>线索 ID：</b>{_esc(finding.get('finding_id'))}　<b>状态：</b>{_esc(finding.get('review_status', '待复核'))}　<b>最终分数：</b>{_esc(finding.get('anomaly_score'))}　<b>规则分数：</b>{_esc(finding.get('rule_score', '-'))}　<b>模型分数：</b>{_esc(finding.get('model_score', '-'))} / {_esc(finding.get('model_threshold', '-'))}　<b>相似度：</b>{_pct(finding.get('similarity'))}　<b>重复片段：</b>{_esc(finding.get('repeated_segment_count', 0))} 段 / {_esc(finding.get('repeated_segment_chars', 0))} 字符</p>
              <p><b>证据页：</b>A {', '.join(map(str, finding.get('evidence_pages_a', []))) or '-'}；B {', '.join(map(str, finding.get('evidence_pages_b', []))) or '-'}</p>
              <ul>{evidence or '<li>未生成结构化证据摘要</li>'}</ul>
              <p class="review">建议：回到原始文件核对投标人主体、页码、共同实体来源和文件形成时间。</p>
            </section>
            """
        )
    if not finding_sections:
        finding_sections.append('<div class="empty">当前阈值下没有待复核异常线索。</div>')
    html_text = f"""<!doctype html>
<html lang="zh-CN"><head><meta charset="utf-8"><title>{_esc(payload.get('title'))}</title>
<style>
body{{font-family:"Microsoft YaHei",Arial,sans-serif;color:#172033;line-height:1.65;margin:0;background:#f5f7fa}}
main{{max-width:1080px;margin:0 auto;background:#fff;padding:42px 54px;min-height:100vh}}
h1{{font-size:28px;margin:0 0 8px}}h2{{font-size:18px;border-bottom:1px solid #d6dde6;padding-bottom:7px;margin-top:28px}}
.meta,.muted{{color:#667085;font-size:13px}}.metrics{{display:grid;grid-template-columns:repeat(3,1fr);gap:10px;margin:22px 0}}
.metric{{border:1px solid #d6dde6;padding:13px;background:#f8fafc}}.metric span{{display:block;color:#667085;font-size:12px}}.metric strong{{font-size:22px}}
.notice{{border-left:4px solid #1769e0;background:#eef5ff;padding:12px 15px;margin:18px 0}}.finding{{border-top:3px solid #d6dde6;padding-top:8px;margin-top:24px}}
.finding h2{{border:0;margin:0 0 8px}}.review{{color:#7a1a12;background:#fff5f3;padding:10px}}.empty{{padding:18px;background:#f8fafc;color:#667085}}
table{{border-collapse:collapse;width:100%}}td,th{{border:1px solid #d6dde6;padding:8px;text-align:left}}th{{background:#f1f5f9}}
@media print{{body{{background:#fff}}main{{padding:0;max-width:none}}.finding{{break-inside:avoid}}}}
</style></head><body><main>
<h1>{_esc(payload.get('title'))}</h1>
<div class="meta">生成时间：{_esc(payload.get('generated_at'))}</div>
<div class="notice">{_esc(payload.get('boundary'))}</div>
<h2>一、分析概览</h2><div class="metrics">{metrics}</div>
<p>分析项目数：{_esc(summary.get('project_count', 0))}；投标文件数：{_esc(summary.get('document_count', 0))}；项目内比较数：{_esc(summary.get('pair_count', 0))}；待复核线索数：{_esc(summary.get('anomaly_count', len(findings)))}</p>
<h2>二、关系图谱摘要</h2><table><tr><th>项目节点</th><th>投标人节点</th><th>文件节点</th><th>实体/关系总数</th></tr><tr><td>{_esc(len([n for n in graph.get('nodes', []) if n.get('type') == 'project']))}</td><td>{_esc(len([n for n in graph.get('nodes', []) if n.get('type') == 'bidder']))}</td><td>{_esc(len([n for n in graph.get('nodes', []) if n.get('type') == 'document']))}</td><td>{_esc(graph.get('node_count', 0))} / {_esc(graph.get('edge_count', 0))}</td></tr></table>
<h2>三、异常线索明细</h2>{''.join(finding_sections)}
 <h2>四、模型和大模型状态</h2><p>模型类型：{_esc(payload.get('model', {}).get('model_type', '未读取'))}；训练比较数：{_esc(model_summary.get('pair_count', payload.get('model', {}).get('training_pair_count', 0)))}；标签状态：{_esc(model_summary.get('label_status', 'unlabeled_unsupervised'))}；规则线索：{_esc(summary.get('rule_anomaly_count', '-'))}；模型触发：{_esc(summary.get('model_triggered_count', '-'))}。</p>
<p>大模型状态：{_esc(llm.get('status', 'skipped'))}；通过本地引用校验：{_esc(llm.get('validated_finding_count', 0))} 条。大模型只用于辅助解释候选线索。</p>
<h2>五、复核建议</h2><ol><li>按照线索 ID、投标人和页码回到原始文件。</li><li>核对共同电话、地址、邮箱、联系人和文件形成时间是否具有业务合理性。</li><li>结合开标记录、报价表、评标材料和外部登记信息进行人工判断。</li><li>将复核结果记录为待复核、排除或需要进一步调查，不能仅凭模型分数定性。</li></ol>
<p class="muted">模型版本：{_esc(payload.get('model', {}).get('schema_version', '未读取'))}；报告数据版本：{_esc(payload.get('analysis', {}).get('summary', {}).get('output_dir', 'local-artifacts'))}</p>
</main></body></html>"""
    return html_text.encode("utf-8")


def _set_cell(cell, text: object, *, bold: bool = False) -> None:
    cell.text = str(text if text is not None else "")
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run.font.size = Pt(9)
            run.bold = bold


def build_docx_report(payload: dict) -> bytes:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    styles = document.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10)
    document.add_heading(payload.get("title", "招采智审异常线索复核报告"), level=0)
    document.add_paragraph(f"生成时间：{payload.get('generated_at', '')}")
    boundary = document.add_paragraph()
    boundary.add_run("结论边界：").bold = True
    boundary.add_run(payload.get("boundary", ""))

    document.add_heading("一、分析概览", level=1)
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    _set_cell(table.rows[0].cells[0], "指标", bold=True)
    _set_cell(table.rows[0].cells[1], "数值", bold=True)
    for label, value in _metric_rows(payload):
        cells = table.add_row().cells
        _set_cell(cells[0], label)
        _set_cell(cells[1], value)

    summary = payload.get("summary", {})
    document.add_paragraph(
        f"分析项目数：{summary.get('project_count', 0)}；投标文件数：{summary.get('document_count', 0)}；"
        f"项目内比较数：{summary.get('pair_count', 0)}；待复核线索数：{summary.get('anomaly_count', 0)}。"
    )

    graph = payload.get("graph", {})
    document.add_heading("二、关系图谱摘要", level=1)
    document.add_paragraph(
        f"图谱包含 {graph.get('node_count', 0)} 个节点和 {graph.get('edge_count', 0)} 条关系，"
        f"覆盖 {graph.get('project_count', 0)} 个项目。关系用于定位证据入口，不代表违规结论。"
    )

    document.add_heading("三、异常线索明细", level=1)
    findings = payload.get("analysis", {}).get("anomalies", [])
    if not findings:
        document.add_paragraph("当前阈值下没有待复核异常线索。")
    for index, finding in enumerate(findings, start=1):
        document.add_heading(f"{index}. {finding.get('bidder_a', '')} 对比 {finding.get('bidder_b', '')}", level=2)
        document.add_paragraph(
            f"线索 ID：{finding.get('finding_id', '')}\n"
            f"项目：{finding.get('project_id', '')}\n"
            f"状态：{finding.get('review_status', '待复核')}；最终分数：{finding.get('anomaly_score', '')}；"
            f"规则分数：{finding.get('rule_score', '-')}；模型分数：{finding.get('model_score', '-')} / {finding.get('model_threshold', '-')}；"
            f"文本相似度：{_pct(finding.get('similarity'))}；重复片段：{finding.get('repeated_segment_count', 0)} 段 / {finding.get('repeated_segment_chars', 0)} 字符\n"
            f"证据页：A {', '.join(map(str, finding.get('evidence_pages_a', []))) or '-'}；"
            f"B {', '.join(map(str, finding.get('evidence_pages_b', []))) or '-'}"
        )
        for item in finding.get("evidence", []):
            document.add_paragraph(str(item), style="List Bullet")
        document.add_paragraph("建议复核：回到原始文件核对投标人主体、页码、共同实体来源和文件形成时间。")

    document.add_heading("四、模型和大模型状态", level=1)
    model = payload.get("model", {})
    model_summary = payload.get("model_summary", {})
    llm = payload.get("llm", {})
    document.add_paragraph(
        f"模型类型：{model.get('model_type', '未读取')}；训练比较数："
        f"{model_summary.get('pair_count', model.get('training_pair_count', 0))}；"
        f"标签状态：{model_summary.get('label_status', 'unlabeled_unsupervised')}。"
    )
    document.add_paragraph(
        f"大模型状态：{llm.get('status', 'skipped')}；通过本地引用校验："
        f"{llm.get('validated_finding_count', 0)} 条。大模型只用于辅助解释候选线索。"
    )

    document.add_heading("五、复核建议", level=1)
    for item in (
        "按照线索 ID、投标人和页码回到原始文件。",
        "核对共同电话、地址、邮箱、联系人和文件形成时间是否具有业务合理性。",
        "结合开标记录、报价表、评标材料和外部登记信息进行人工判断。",
        "将复核结果记录为待复核、排除或需要进一步调查，不能仅凭模型分数定性。",
    ):
        document.add_paragraph(item, style="List Number")

    output = __import__("io").BytesIO()
    document.save(output)
    return output.getvalue()
