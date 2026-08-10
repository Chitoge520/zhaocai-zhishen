from __future__ import annotations

import csv
import json
import tempfile
import unittest
from pathlib import Path

from docx import Document

from zhaocai_zhishen.document_pipeline import ExtractedPage, clean_extracted_pages, load_samples, normalize_text, prepare_dataset


class DocumentPipelineTests(unittest.TestCase):
    def test_normalize_text(self) -> None:
        self.assertEqual(normalize_text("  第一行  \n\n 第二\t行 "), "第一行\n第二 行")

    def test_prepare_docx_from_manifest(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            root = Path(temp)
            project_dir = root / "项目001" / "投标文件"
            project_dir.mkdir(parents=True)
            doc_path = project_dir / "A公司.docx"
            doc = Document()
            doc.add_paragraph("测试项目投标文件")
            doc.add_paragraph("投标人：A公司")
            doc.save(doc_path)

            with (root / "samples.csv").open("w", encoding="utf-8", newline="") as handle:
                writer = csv.DictWriter(handle, fieldnames=["sample_id", "project_name", "bidder_candidate", "bid_file", "source_inner_path"])
                writer.writeheader()
                writer.writerow({
                    "sample_id": "sample-001",
                    "project_name": "项目001",
                    "bidder_candidate": "A公司",
                    "bid_file": "项目001\\投标文件\\A公司.docx",
                    "source_inner_path": "原始/A公司.docx",
                })

            self.assertEqual(len(load_samples(root)), 1)
            output = root / "processed"
            summary = prepare_dataset(root, output, ocr_mode="off")
            self.assertEqual(summary["success_count"], 1)
            self.assertEqual(summary["ocr_engine"], "off")
            row = json.loads((output / "documents.jsonl").read_text(encoding="utf-8").splitlines()[0])
            self.assertEqual(row["document_id"], "sample-001")
            self.assertIn("测试项目", row["text"])
            self.assertEqual(row["parse_status"], "success")
            cached_summary = prepare_dataset(root, output, ocr_mode="off")
            self.assertEqual(cached_summary["cached_count"], 1)

    def test_clean_extracted_pages_preserves_raw_and_removes_ocr_repetition(self) -> None:
        pages = [
            ExtractedPage(1, "项目投标文件\n有效条款甲\n第1页", 0, "paddleocr"),
            ExtractedPage(2, "项目投标文件\n" + "有效条款乙\n" * 12 + "第2页", 0, "paddleocr"),
            ExtractedPage(3, "项目投标文件\n" + "有效条款乙\n" * 12 + "第2页", 0, "paddleocr"),
        ]

        stats = clean_extracted_pages(pages)

        self.assertIn("项目投标文件", pages[0].text)
        self.assertIn("项目投标文件", pages[0].cleaned_text)
        self.assertNotIn("项目投标文件", pages[1].cleaned_text)
        self.assertEqual(pages[2].cleaned_text, "")
        self.assertEqual(stats["duplicate_page_count"], 1)
        self.assertGreater(stats["removed_redundant_line_count"], 0)


if __name__ == "__main__":
    unittest.main()
